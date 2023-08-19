import os
from docx import Document

phonebook_file = "phonebook.docx"
column_names = ['Фамилия', 'Имя', 'Отчество', 'Организация', 'Рабочий тел.', 'Личный тел.']

def check_phonebook():
    if os.path.exists(phonebook_file):
        doc = Document(phonebook_file)
        data = []
        for row in doc.tables[0].rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        return doc, data  # Возвращаем doc вместе с данными
    else:
        print("404")

def print_data():
    _, data = check_phonebook()  # Игнорируем возвращаемое значение doc
    for row in data:
        print(row)
    print("\n")

def enter_new_record():
    new_data = []
    for col, column_name in enumerate(column_names):
        value = input(f"Введите значение для столбца {column_name}: ")
        new_data.append(value)

    doc, _ = check_phonebook()  # Получаем doc для добавления записи
    table = doc.tables[0]
    new_row = table.add_row().cells
    for col_idx, cell_text in enumerate(new_data):
        new_row[col_idx].text = cell_text
    print()
    doc.save(phonebook_file)  # Сохранение обновленного файла



def edit_record():
    doc, data = check_phonebook()
    table = doc.tables[0]

    print("Список доступных записей:")
    for idx, row in enumerate(data):
        print(f"{idx + 1}. {', '.join(row)}")

    record_number = int(input("Введите номер записи, которую хотите отредактировать: ")) - 1

    if 0 <= record_number < len(data):
        edited_data = []
        for col, value in enumerate(data[record_number]):
            new_value = input(f"Введите новое значение для столбца '{column_names[col]}': ")
            edited_data.append(new_value)

        for cell, new_value in zip(table.rows[record_number+1].cells, edited_data):
            cell.text = new_value

        doc.save(phonebook_file)
        print("Запись успешно отредактирована.")
    else:
        print("Некорректный номер записи.")

def single_search():
    _, data = check_phonebook()

    search_column = input("Введите название столбца для поиска (например, 'Фамилия'): ")
    search_value = input(f"Введите значение для поиска в столбце '{search_column}': ")

    found_records = []
    for row in data:
        if search_value in row[column_names.index(search_column)]:
            found_records.append(row)

    if found_records:
        print("Найденные записи:")
        for idx, row in enumerate(found_records):
            print(f"{idx + 1}. {', '.join(row)}")
    else:
        print("Записи не найдены.")

def mono_search():
    _, data = check_phonebook()

    search_conditions = []
    while True:
        search_column = input(
            "Введите название столбца для поиска (например, 'Фамилия', или введите 'готово' для завершения): ")
        if search_column.lower() == 'готово':
            break
        search_value = input(f"Введите значение для поиска в столбце '{search_column}': ")
        search_conditions.append((search_column, search_value))

    found_records = []
    for row in data:
        matched = True
        for search_column, search_value in search_conditions:
            if search_value not in row[column_names.index(search_column)]:
                matched = False
                break
        if matched:
            found_records.append(row)

    if found_records:
        print("Найденные записи:")
        for idx, row in enumerate(found_records):
            print(f"{idx + 1}. {', '.join(row)}")
    else:
        print("Записи не найдены.")


def main():
    while True:
        print(" 1. Посмотреть записи в телефонной книге ")
        print(" 2. Добавить новую запись в справочник ")
        print(" 3. Редактировать запись в справочнике ")
        print(" 4. Найти запись по одному значению ")
        print(" 5. Найти запись по нескольким значениям ")
        print(" 6. Выйти \n")

        choice = input()

        if choice == '1':
            print_data()
        if choice == '2':
            enter_new_record()
        if choice == '3':
            edit_record()
        if choice == '4':
            single_search()
        if choice == '5':
            mono_search()
        if choice == '6':
            break
        if int(choice) > 6 or int(choice) < 1:
            print("Такой вариант ответа не предусмотрен")

if __name__ == '__main__':
    main()
